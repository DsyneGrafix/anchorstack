from flask import Flask, render_template_string, request, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)

HTML_FORM = """
<!doctype html>
<title>Watchman Chapter Generator</title>
<h2>Generate Watchman Chapters</h2>
<form method=post>
  Book Title: <input type=text name=title value="The Watchman"><br><br>
  Chapter Base Title: <input type=text name=base value="The Watchman’s Call"><br><br>
  <input type=submit value=Generate>
</form>
"""

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        title = request.form['title']
        base = request.form['base']
        
        doc = Document()
        doc.add_heading(title, 0)

        for i in range(1, 101):
            doc.add_heading(f'Chapter {i}: {base} {i}', level=1)
            doc.add_paragraph(
                ("This is a placeholder for Chapter {} of '{}'. ".format(i, title) * 25)[:500]
            )
            if i % 2 == 0:
                doc.add_paragraph("[Insert fractal image here]")

        doc.add_page_break()
        doc.add_heading('Thank You', level=1)
        doc.add_paragraph("Thanks for reading. Stay tuned for the next volume!")

        # Save to memory
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(buffer, as_attachment=True, download_name=f"{title.replace(' ', '_')}.docx")

    return render_template_string(HTML_FORM)

if __name__ == '__main__':
    app.run(debug=True)

